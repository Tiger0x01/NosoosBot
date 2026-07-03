import os
import datetime
import re
import textwrap
from docx import Document
from fpdf import FPDF
import arabic_reshaper
from bidi.algorithm import get_display
from config import FONTS_FOLDER, BASE_DOWNLOAD_FOLDER

def sanitize_filename(title):
    """تنظيف اسم الملف من الرموز الممنوعة"""
    return re.sub(r'[\\/*?:"<>|]', "", title)[:40].strip()

def clean_markdown(text):
    """تنظيف النص من كل علامات الماركداون المزعجة"""
    text = re.sub(r'[*_#`]', '', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text) # إزالة الروابط وترك النص
    text = re.sub(r'^>+', '', text, flags=re.MULTILINE) # إزالة الاقتباسات
    return text.strip()

def get_daily_folder():
    """إنشاء مجلد يومي لتنظيم الملفات"""
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    folder_path = os.path.join(BASE_DOWNLOAD_FOLDER, current_date)
    os.makedirs(folder_path, exist_ok=True)
    return folder_path

def get_font_path():
    """البحث عن ملف الخط بأكثر من صيغة لمنع أخطاء الـ Case-Sensitivity"""
    for font_name in ["arial.ttf", "ARIAL.TTF", "Arial.ttf"]:
        font_path = os.path.join(FONTS_FOLDER, font_name)
        if os.path.exists(font_path):
            return font_path
    # لو الخط مش موجود، نوقف العملية برسالة واضحة
    raise FileNotFoundError("❌ ملف خط 'arial.ttf' غير موجود في مجلد 'fonts'. يرجى إضافته لدعم اللغة العربية.")

def generate_files(text, base_filename, video_title="Video Transcript", video_url="", language=""):
    folder = get_daily_folder()
    safe_title = sanitize_filename(video_title)
    txt_path = os.path.join(folder, f"{safe_title}_{base_filename}.txt")
    pdf_path = os.path.join(folder, f"{safe_title}_{base_filename}.pdf")
    docx_path = os.path.join(folder, f"{safe_title}_{base_filename}.docx")

    header = f"Title: {video_title}\nVideo URL: {video_url}\nLanguage: {language}\nDate: {datetime.datetime.now().strftime('%Y-%m-%d')}\n{'-'*40}\n\n"
    full_text = header + text

    # ================= 1. TXT =================
    with open(txt_path, "w", encoding="utf-8") as f: 
        f.write(full_text)

    # ================= 2. DOCX =================
    doc = Document()
    doc.add_heading(video_title, level=1)
    doc.add_paragraph(f"Video URL: {video_url}\nLanguage: {language}\n{'-'*40}")
    doc.add_paragraph(text)
    doc.save(docx_path)

    # ================= 3. PDF =================
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # جلب مسار الخط الآمن
    font_path = get_font_path()
    
    try: pdf.add_font("Arabic", "", font_path, uni=True)
    except: pdf.add_font("Arabic", "", font_path)
    pdf.set_font("Arabic", size=12)

    max_w = pdf.w - 30 # عرض الصفحة ناقص الهوامش
    
    for line in full_text.splitlines():
        if not line.strip():
            pdf.ln(5)
            continue
        
        words = line.split()
        current_line = ""
        for word in words:
            test_line = f"{current_line} {word}".strip()
            if pdf.get_string_width(test_line) <= max_w:
                current_line = test_line
            else:
                if language.startswith("ar"):
                    pdf.cell(w=0, h=8, txt=get_display(arabic_reshaper.reshape(current_line)), ln=1, align="R")
                else:
                    pdf.cell(w=0, h=8, txt=current_line, ln=1, align="L")
                current_line = word
        
        if current_line:
            if language.startswith("ar"):
                pdf.cell(w=0, h=8, txt=get_display(arabic_reshaper.reshape(current_line)), ln=1, align="R")
            else:
                pdf.cell(w=0, h=8, txt=current_line, ln=1, align="L")

    pdf.output(pdf_path)
    return txt_path, pdf_path, docx_path

def generate_summary_pdf(summary_text, video_title, user_id):
    folder = get_daily_folder()
    safe_title = sanitize_filename(video_title)
    pdf_path = os.path.join(folder, f"Summary_{safe_title}_{user_id}.pdf")
    
    clean_summary = clean_markdown(summary_text) 
    full_text = f"ملخص فيديو: {video_title}\n{'-'*40}\n\n" + clean_summary

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # جلب مسار الخط الآمن
    font_path = get_font_path()
    
    try: pdf.add_font("Arabic", "", font_path, uni=True)
    except: pdf.add_font("Arabic", "", font_path)
    pdf.set_font("Arabic", size=12)

    max_w = pdf.w - 30
    for line in full_text.splitlines():
        if not line.strip():
            pdf.ln(5)
            continue
        
        words = line.split()
        current_line = ""
        for word in words:
            test_line = f"{current_line} {word}".strip()
            if pdf.get_string_width(test_line) <= max_w: 
                current_line = test_line
            else:
                # نستخدم اللغة العربية إجبارياً هنا لأن الملخص دائماً بالعربي
                pdf.cell(w=0, h=8, txt=get_display(arabic_reshaper.reshape(current_line)), ln=1, align="R")
                current_line = word
                
        if current_line:
            pdf.cell(w=0, h=8, txt=get_display(arabic_reshaper.reshape(current_line)), ln=1, align="R")

    pdf.output(pdf_path)
    return pdf_path